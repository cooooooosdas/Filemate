"""Word 解析。"""

from __future__ import annotations

import logging
from pathlib import Path

from ..parsers import register

logger = logging.getLogger(__name__)


class WordParser:
    """Word 解析器。

    - .docx → python-docx
    - .doc  → 旧格式，python-docx 不支持，返回带 error 的结构
    """

    def parse(self, path: str | Path) -> dict:
        p = Path(path)
        suffix = p.suffix.lstrip(".").lower()

        if suffix == "doc":
            return {
                "raw_text": "",
                "metadata": {"suffix": "doc", "note": "旧 .doc 格式暂不支持，请另存为 .docx"},
                "error": "旧 .doc 格式暂不支持，请另存为 .docx 后重试",
            }

        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError(
                "python-docx 未安装。运行: pip install python-docx"
            ) from exc

        doc = Document(str(p))
        paragraphs: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        raw = "\n".join(paragraphs)
        logger.debug("Word 解析: %s → %d 字", p.name, len(raw))
        return {
            "raw_text": raw,
            "metadata": {
                "suffix": "docx",
                "paragraphs": len(paragraphs),
                "tables": len(doc.tables),
            },
        }


register("docx", WordParser)
register("DOCX", WordParser)
